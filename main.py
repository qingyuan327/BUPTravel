from flask import Flask, render_template, request, redirect, url_for, session, jsonify, flash
from test import con_my_sql,con_my_sql_with_params
from flask_sqlalchemy import SQLAlchemy
from db_utils import con_my_sql, con_my_sql_with_params
import zlib
import pymysql
pymysql.install_as_MySQLdb()
from whoosh.index import create_in, open_dir
from whoosh.fields import Schema, TEXT, ID
from whoosh.qparser import QueryParser
import numpy as np
from sklearn.preprocessing import MinMaxScaler
import os
from io import BytesIO
from datetime import datetime
from functools import wraps
from flask import send_file
import gzip
from docx import Document  # 导入 python-docx 库
from docx.shared import Inches
import cv2
import numpy as np
from PIL import Image, ImageDraw, ImageFont
import tempfile
import requests
import json
from zhipuai import ZhipuAI
from werkzeug.utils import secure_filename
from threading import Thread
import time
import logging
from threading import Lock
from flask_cors import CORS
import base64
import heapq
import math
import heapq  # 添加这行导入
from collections import defaultdict

#创建 Flask 应用实例
app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] ='mysql+pymysql://root:mahaijuan0511@localhost/demo01'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.secret_key = 'your_secret_key'  # 添加 secret_key 用于 flash 消息
# app.config['SQLALCHEMY_DATABASE_URI'] ='sqlite:///your_database.db'
db = SQLAlchemy(app)
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
CORS(app, resources={r"/api/*": {"origins": "*"}})  # 允许所有来源访问 /api 路径

# 智谱清影API配置（需替换为你的实际密钥）
ZHI_PU_API_KEY = "0014ce0d9cd446d79cb2f1ee7737b050.T5pA1tcnXEsxNn3c"
ZHI_PU_API_URL = "https://open.bigmodel.cn/api/paas/v4/videos/generations"
ZHI_PU_RESULT_URL = "https://open.bigmodel.cn/api/paas/v4/async-result/{}"

# 任务状态跟踪
task_status = {}
status_lock = Lock()  # 添加锁

db_config = {
    'host': 'localhost',
    'user': 'root',
    'password': 'mahaijuan0511',
    'database': 'demo01'
}


# 定义用户模型
class LoginUser(db.Model):
    __tablename__ = 'login_user'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(45), unique=True, nullable=False)
    password = db.Column(db.String(45), nullable=False)


# 定义景点模型
class Attraction(db.Model):
    __tablename__ = 'attraction'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(255))
    heat = db.Column(db.Integer)
    image_path = db.Column(db.String(255))
    score = db.Column(db.Integer)
    recommend = db.Column(db.Float)
    type = db.Column(db.String(50))

# 定义景点评分模型
class AttractionRatings(db.Model):
    __tablename__ = 'attraction_ratings'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer)
    attraction_id = db.Column(db.Integer)
    rating = db.Column(db.Float)
    __table_args__ = (
        db.UniqueConstraint('user_id', 'attraction_id', name='_user_attraction_uc'),
    )



#定义道路信息模型
class RoadLinks(db.Model):
    __tablename__ = 'roadlinks'
    road_id = db.Column(db.String(10), primary_key=True)  # 联合主键或唯一约束
    from_node = db.Column(db.String(10), nullable=False)
    to_node = db.Column(db.String(10), nullable=False)
    length = db.Column(db.Integer, nullable=False)
    __table_args__ = (
        db.UniqueConstraint('road_id', 'from_node', 'to_node', name='_roadlink_uc'),
    )

#定义道路节点模型
class Node(db.Model):
    __tablename__ = 'nodes'  # 对应数据表名
    node_id = db.Column(db.String(10), primary_key=True)  # 节点ID，如 P001、S001
    name = db.Column(db.String(255), nullable=False)       # 节点名称，如 "校训石"、"物美超市"
    category = db.Column(db.Integer, nullable=False)        # 类别（0=place，1=service）
    type = db.Column(db.String(50), nullable=False)         # 类型，如 "景点"、"超市"

#定义服务设施模型
class Service(db.Model):
    __tablename__ = 'service'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(255))
    type = db.Column(db.String(255))

# 定义一个装饰器，用于检查用户是否登录
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'username' not in session:
            return redirect(url_for('index_login'))
        return f(*args, **kwargs)
    return decorated_function

# 模拟 before_first_request 的功能
first_request = True
@app.before_request
def clear_session_on_first_request():
    global first_request
    if first_request:
        session.clear()
        first_request = False

# 定义旅游日记模型
class TravelDiary(db.Model):
    __tablename__ = 'travel_diaries'
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(255))
    content = db.Column(db.Text)
    image_path = db.Column(db.String(255))
    video_path = db.Column(db.String(255))
    rating = db.Column(db.Float, default=0)
    views = db.Column(db.Integer, default=0)  # 将 views 当作总点赞数量
    destination = db.Column(db.String(255))
    comments = db.relationship('Comment', backref='diary', lazy=True)  # 添加关系

    __searchable__ = ['title', 'content', 'destination']  # 指定需要搜索的字段

# 定义点赞模型
class Like(db.Model):
    __tablename__ = 'likes'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(45), db.ForeignKey('login_user.username'))
    diary_id = db.Column(db.Integer, db.ForeignKey('travel_diaries.id'))

# 定义评论模型
class Comment(db.Model):
    __tablename__ = 'comments'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(45), db.ForeignKey('login_user.username'))
    diary_id = db.Column(db.Integer, db.ForeignKey('travel_diaries.id'))
    comment = db.Column(db.Text)
    comment_time = db.Column(db.DateTime, default=datetime.utcnow)  # 添加评论时间字段

# 定义评分模型
class Rating(db.Model):
    __tablename__ = 'ratings'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(45), db.ForeignKey('login_user.username'))
    diary_id = db.Column(db.Integer, db.ForeignKey('travel_diaries.id'))
    rating = db.Column(db.Float, nullable=False)

# 定义场所模型
class Place(db.Model):
    __tablename__ = 'place'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(255))
    heat = db.Column(db.Integer)
    intro = db.Column(db.Text)
    image_path = db.Column(db.String(255))
    type = db.Column(db.String(50))  # 对应 place.txt 中的 type
# 定义 美食 模型
class Food(db.Model):
    __tablename__ = 'food'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100))  # 调整为100字符
    heat = db.Column(db.Integer)
    intro = db.Column(db.String(255))  # 使用String替代Text，长度255
    image_path = db.Column(db.String(255))
    cuisine = db.Column(db.String(255))  # 新增字段，类型varchar(255)


# 景区和校园地点模型（包含地理坐标）
class Location(db.Model):
    __tablename__ = 'locations'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(255), nullable=False)
    type = db.Column(db.String(50), nullable=False)  # 景区、建筑物、服务设施等
    latitude = db.Column(db.Float, nullable=False)   # 纬度
    longitude = db.Column(db.Float, nullable=False)  # 经度
    description = db.Column(db.Text)
    image_path = db.Column(db.String(255))
    heat = db.Column(db.Integer, default=0)
    score = db.Column(db.Float, default=0.0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

# 路径节点模型
class PathNode(db.Model):
    __tablename__ = 'path_nodes'
    id = db.Column(db.Integer, primary_key=True)
    location_id = db.Column(db.Integer, db.ForeignKey('locations.id'), nullable=False)
    node_type = db.Column(db.String(50), default='normal')  # normal, entrance, exit, junction
    is_accessible = db.Column(db.Boolean, default=True)

# 路径边模型（连接两个节点的路径）
class PathEdge(db.Model):
    __tablename__ = 'path_edges'
    id = db.Column(db.Integer, primary_key=True)
    from_node_id = db.Column(db.Integer, db.ForeignKey('path_nodes.id'), nullable=False)
    to_node_id = db.Column(db.Integer, db.ForeignKey('path_nodes.id'), nullable=False)
    distance = db.Column(db.Float, nullable=False)  # 距离（米）
    travel_time_walk = db.Column(db.Float, nullable=False)  # 步行时间（分钟）
    travel_time_bike = db.Column(db.Float, default=0)  # 骑行时间（分钟）
    travel_time_bus = db.Column(db.Float, default=0)   # 公交时间（分钟）
    path_type = db.Column(db.String(50), default='walkway')  # walkway, road, stairs等
    is_bidirectional = db.Column(db.Boolean, default=True)

# 旅游路线模型
class TourRoute(db.Model):
    __tablename__ = 'tour_routes'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.String(45), db.ForeignKey('login_user.username'))
    route_name = db.Column(db.String(255), nullable=False)
    start_location_id = db.Column(db.Integer, db.ForeignKey('locations.id'))
    end_location_id = db.Column(db.Integer, db.ForeignKey('locations.id'))
    waypoints = db.Column(db.Text)  # JSON格式存储途经点
    total_distance = db.Column(db.Float)
    total_time = db.Column(db.Float)
    transport_mode = db.Column(db.String(50), default='walk')  # walk, bike, bus, mixed
    route_data = db.Column(db.Text)  # JSON格式存储完整路径数据
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

# 定义用户反馈模型
class Feedback(db.Model):
    __tablename__ = 'feedback'
    id = db.Column(db.Integer, primary_key=True)
    mailbox = db.Column(db.String(255))
    content = db.Column(db.String(255))
    
#搜索时关键词高亮显示
def highlight_keyword(text, keyword):
    if keyword:
        highlighted = text.replace(keyword, f'<span class="highlight">{keyword}</span>')
        return highlighted
    return text

#计算景点热度
def calculate_attraction_recommend():
    attractions = Attraction.query.all()
    for attraction in attractions:
        if attraction.heat is not None and attraction.score is not None:
            attraction.recommend = (0.4 * attraction.heat) + (0.6 * attraction.score)
    db.session.commit()

#景点评分系统
def update_attraction_scores():
    attractions = Attraction.query.all()
    for attraction in attractions:
        ratings = AttractionRatings.query.filter_by(attraction_id=attraction.id).all()
        if ratings:
            total_rating = sum([r.rating for r in ratings])
            average_rating = total_rating / len(ratings)
            # 进位处理
            attraction.score = math.ceil(average_rating)
        else:
            attraction.score = 0
    db.session.commit()


#核心算法：不经过完全排序排好前10的景点
def get_top_10_attractions(search_query, sort_by, category):
    # 先获取所有景点数据
    all_attractions = Attraction.query.all()

    if not all_attractions:
        return []

    # 使用线性查找筛选符合条件的景点
    attractions = []
    for attraction in all_attractions:
        match_search = True
        match_category = True

        # 检查搜索条件
        if search_query:
            name = attraction.name if attraction.name else ''
            match_search = search_query in name

        # 检查类别过滤条件
        if category:
            attraction_type = attraction.type if attraction.type else ''
            match_category = attraction_type == category

        # 如果同时满足搜索和类别条件，则添加到筛选结果中
        if match_search and match_category:
            attractions.append(attraction)

    # 获取当前登录用户的ID
    user_id = None
    if 'username' in session:
        user = LoginUser.query.filter_by(username=session['username']).first()
        if user:
            user_id = user.id

    # 定义排序规则
    if sort_by == 'heat':
        def key_func(attraction):
            return attraction.heat
    elif sort_by == 'rating':
        def key_func(attraction):
            return attraction.score
    elif sort_by == 'recommend':
        def key_func(attraction):
            recommend_score = attraction.recommend
            if user_id:
                user_rating = AttractionRatings.query.filter_by(user_id=user_id, attraction_id=attraction.id).first()
                if user_rating and user_rating.rating >= 3:
                    # 给评分在三颗星以上的景点增加一定权值，这里假设增加权值为 2
                    recommend_score += 2
            return recommend_score
    else:
        def key_func(attraction):
            recommend_score = attraction.recommend
            if user_id:
                user_rating = AttractionRatings.query.filter_by(user_id=user_id, attraction_id=attraction.id).first()
                if user_rating and user_rating.rating >= 3:
                    # 给评分在三颗星以上的景点增加一定权值，这里假设增加权值为 2
                    recommend_score += 2
            return recommend_score

    # 手动实现的堆排序算法

    # 如果景点列表 attractions 为空，直接返回空列表
    if not attractions:
        return []

    # 创建一个最小堆（用列表表示）
    heap = []
    k = 10

    for attraction in attractions:
        score = key_func(attraction)
        # 堆未满时直接添加
        if len(heap) < k:
            heap.append((score, attraction))
            # 上浮调整堆
            _sift_up(heap, len(heap) - 1)
        else:
            # 堆已满，比较当前元素与堆顶元素
            if score > heap[0][0]:
                # 替换堆顶元素并下沉调整
                heap[0] = (score, attraction)
                _sift_down(heap, 0, len(heap))

    # 对堆中的元素进行排序（从大到小）
    result = [item[1] for item in heap]

    # 快速排序实现（降序）
    def quick_sort_descending(arr, key_func):
        if len(arr) <= 1:
            return arr
        # 选择基准元素（这里选中间元素，也可以随机选择优化）
        pivot = arr[len(arr) // 2]
        # 划分左右子数组
        left = [x for x in arr if key_func(x) > key_func(pivot)]
        middle = [x for x in arr if key_func(x) == key_func(pivot)]
        right = [x for x in arr if key_func(x) < key_func(pivot)]
        # 递归排序并合并结果
        return quick_sort_descending(left, key_func) + middle + quick_sort_descending(right, key_func)

    # 执行快速排序
    result = quick_sort_descending(result, key_func)
    return result


# 堆调整辅助函数：上浮元素以维护最小堆性质
def _sift_up(heap, idx):
    while idx > 0:
        parent_idx = (idx - 1) // 2
        if heap[idx][0] >= heap[parent_idx][0]:
            break
        # 交换元素
        temp = heap[idx]
        heap[idx] = heap[parent_idx]
        heap[parent_idx] = temp
        idx = parent_idx


# 堆调整辅助函数：下沉元素以维护最小堆性质
def _sift_down(heap, idx, heap_size):
    while True:
        left = 2 * idx + 1
        right = 2 * idx + 2
        smallest = idx

        if left < heap_size and heap[left][0] < heap[smallest][0]:
            smallest = left
        if right < heap_size and heap[right][0] < heap[smallest][0]:
            smallest = right

        if smallest == idx:
            break

        # 交换元素
        temp = heap[idx]
        heap[idx] = heap[smallest]
        heap[smallest] = temp
        idx = smallest

#显示场所
def get_top_10_places(search_query, sort_by, category):
    # 先获取所有场所数据

    base_sql = "SELECT * FROM place"
    cursor = con_my_sql_with_params(base_sql, [])
    all_places = cursor.fetchall()

    if all_places is None:
        return []

    # 使用线性查找筛选符合条件的场所
    places = []
    for place in all_places:
        match_search = True
        match_category = True

        # 检查搜索条件
        if search_query:
            name = place.get('name', '')
            intro = place.get('intro', '')
            # 检查 intro 是否为 None，若为 None 则设为空字符串(没有这段会出错)
            if intro is None:
                intro = ''
            match_search = search_query in name or search_query in intro

        # 检查类别过滤条件
        if category:
            place_type = place.get('type', '')
            match_category = place_type == category

        # 如果同时满足搜索和类别条件，则添加到筛选结果中
        if match_search and match_category:
            places.append(place)

    def key_func(place):
        return place['heat']
    # 对场所进行排序
    places.sort(key=key_func, reverse=True)

    return places

app.jinja_env.filters['highlight_keyword'] = highlight_keyword

# 假设这里有一个 AIGC 函数的占位，实际使用时需要导入相应库和模型
def generate_travel_animation(image_path, text_description):
    # 这里应该是实际生成动画的代码，目前为空
    return "动画生成结果的占位符"

#搜索日记
def create_whoosh_index():
    schema = Schema(id=ID(stored=True), title=TEXT, content=TEXT, destination=TEXT)
    try:
        os.mkdir("indexdir")
    except:
        pass
    create_in("indexdir", schema)

def add_to_whoosh_index(diary_id, title, content, destination):
    ix = open_dir("indexdir")
    writer = ix.writer()
    writer.add_document(id=str(diary_id), title=title, content=content, destination=destination)
    writer.commit()

def search_whoosh_index(query_str):
    ix = open_dir("indexdir")
    with ix.searcher() as searcher:
        query = QueryParser("content", ix.schema).parse(query_str)
        results = searcher.search(query)
        return [int(result['id']) for result in results]

def highlight_keyword(text, keyword):
    if keyword:
        highlighted = text.replace(keyword, f'<span class="highlight">{keyword}</span>')
        return highlighted
    return text

app.jinja_env.filters['highlight_keyword'] = highlight_keyword

#*************************设置路由*********************************
@app.before_request
def clear_session_on_first_request():
    global first_request
    if first_request:
        session.clear()
        first_request = False


@app.route("/")
def index_login():
    return render_template('login.html')

@app.route("/register")
def index_register():
    return render_template('register.html')

@app.route("/about")
@login_required
def index_about():
    return render_template('about.html')

@app.route("/area")
@login_required
def index_area():
    return render_template('area.html')

@app.route("/cultural")
@login_required
def index_cultural():
    return render_template('cultural.html')

@app.route("/food")
@login_required
def index_food():
    return render_template('food.html')

@app.route('/video_generator')
@login_required
def video_generator():
    return render_template('video_create.html')

@app.route("/food111")
@login_required
def index_food111():
    search_query = request.args.get('q', '').strip()  # 获取搜索词
    cuisine = request.args.get('cuisine', '')  # 获取菜系

    # 构建基础查询
    base_sql = "SELECT * FROM food"

    # 添加搜索条件
    search_conditions = []
    params = []
    if search_query:
        search_conditions.append("(name LIKE %s OR intro LIKE %s)")
        params.extend([f'%{search_query}%', f'%{search_query}%'])

    # 添加菜系过滤条件
    if cuisine:
        search_conditions.append("cuisine = %s")
        params.append(cuisine)

    if search_conditions:
        search_sql = " WHERE " + " AND ".join(search_conditions)
    else:
        search_sql = ""

    # 执行查询
    cursor = con_my_sql_with_params(base_sql + search_sql, params)
    foods = cursor.fetchall()

    # 按热度排序
    def key_func(food):
        return food['heat']

    # 手动实现的堆排序算法
    def heap_sort(arr, key_func):
        def heapify(arr, n, i, key_func):
            largest = i
            left = 2 * i + 1
            right = 2 * i + 2

            if left < n and key_func(arr[left]) > key_func(arr[largest]):
                largest = left

            if right < n and key_func(arr[right]) > key_func(arr[largest]):
                largest = right

            if largest != i:
                arr[i], arr[largest] = arr[largest], arr[i]
                heapify(arr, n, largest, key_func)

        n = len(arr)

        # 构建最大堆
        for i in range(n // 2 - 1, -1, -1):
            heapify(arr, n, i, key_func)

        # 一个个交换元素
        for i in range(n - 1, 0, -1):
            arr[0], arr[i] = arr[i], arr[0]
            heapify(arr, i, 0, key_func)

        return arr

    sorted_foods = heap_sort(foods, key_func)
    top_10_foods = sorted_foods[-10:][::-1]

    keyword = request.args.get('q', '')
    return render_template('food111.html', foods=top_10_foods, keyword=keyword, cuisine=cuisine)

# 增加景点热度
@app.route("/increment_attraction_heat/<string:id>")
@login_required
def increment_attraction_heat(id):
    attraction = Attraction.query.get(id)
    if attraction:
        attraction.heat = (attraction.heat or 0) + 1
        db.session.commit()
    return redirect(url_for('place2', id=id))  # 重定向到place2页面

# 景点评分
@app.route('/rate_attraction', methods=['POST'])
@login_required
def rate_attraction():
    data = request.get_json()
    attraction_id = data.get('attraction_id')
    rating = data.get('rating')

    # 获取用户的 user_id
    user = LoginUser.query.filter_by(username=session['username']).first()
    user_id = user.id if user else None

    # 检查用户是否已有评分记录
    existing_rating = AttractionRatings.query.filter_by(user_id=user_id, attraction_id=attraction_id).first()

    if existing_rating:
        # 更新已有评分记录
        existing_rating.rating = rating
    else:
        # 新增评分记录
        new_rating = AttractionRatings(user_id=user_id, attraction_id=attraction_id, rating=rating)
        db.session.add(new_rating)

    try:
        db.session.commit()
        # 更新景点的平均评分
        update_attraction_scores()# 这里需要实现更新平均评分的逻辑
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@app.route("/index")
@login_required
def index():
    return render_template('index.html')

# 景点查询页面
@app.route("/attraction")
@login_required
# 获取前10个景点
def index_attraction():
    search_query = request.args.get('q', '').strip()  # 获取搜索词
    sort_by = request.args.get('sort_by', 'recommend')  # 获取排序方式，默认为综合排序
    category = request.args.get('category', '')  # 获取类别
    calculate_attraction_recommend() #这里需要实现计算综合评分的逻辑
    # 获取前10个景点
    attractions = get_top_10_attractions(search_query, sort_by, category)
    keyword = request.args.get('q', '')
    return render_template('attraction.html', attractions=attractions, keyword=keyword, sort_by=sort_by, category=category)

# 景点详情页面
@app.route("/place2/<string:id>")
@login_required
def place2(id):
    query_sql = "SELECT * FROM attraction WHERE id = %s"
    cursor = con_my_sql_with_params(query_sql, (id,))
    attraction_info = cursor.fetchone()
    if not attraction_info:
        return render_template("404.html")
    if attraction_info['image_path']:
        image_path = attraction_info['image_path'].replace('\\', '/').lstrip('/')
    else :
        image_path = 'images/empty2.jpg'.replace('\\', '/').lstrip('/')

    # 获取用户的 user_id
    user = LoginUser.query.filter_by(username=session['username']).first()
    user_id = user.id if user else None

    # 获取用户对该景点的已有评分
    rating_query = "SELECT rating FROM attraction_ratings WHERE user_id = %s AND attraction_id = %s"
    rating_cursor = con_my_sql_with_params(rating_query, (user_id, int(id)))
    user_rating = rating_cursor.fetchone()
    if user_rating:
        user_rating = user_rating['rating']
    else:
        user_rating = 0

    # 获取用户选择的类别
    category = request.args.get('category', '')
    category_input = request.args.get('category_input', '').strip()
    search_query = request.args.get('q', '').strip()  # 获取搜索词
    sort_by = request.args.get('sort_by', 'recommend')  # 获取排序方式，默认为综合排序

    # 查询附近的场所
    nearby_attractions = []
    nearby_attractions = nearby_attractions[:10]

    places = get_top_10_places(search_query, sort_by, category)

    return render_template("place2.html", attraction_info=attraction_info,
                           image_path=image_path, user_rating=user_rating,
                           nearby_attractions=nearby_attractions, category=category, category_input=category_input,
                           places=places, keyword=search_query, sort_by=sort_by)

@app.route("/traffic")
@login_required
def index_traffic():
    return render_template('video_create.html')

@app.route("/travel")
@login_required
def index_travel():
    search_query = request.args.get('q', '').strip()  # 获取搜索词
    sort_by = request.args.get('sort_by', 'recommend')  # 获取排序方式，默认为综合排序
    category = request.args.get('category', '')  # 获取类别
    places = get_top_10_places(search_query, sort_by, category)
    keyword = request.args.get('q', '')
    return render_template('travel.html', places=places, keyword=keyword, sort_by=sort_by, category=category)
@app.route("/place3/<string:id>")

@login_required
def place3(id):
    # 获取当前场所信息（避免查询不存在的字段）
    place_info = db.session.query(
        Place.id,
        Place.name,
        Place.heat,
        Place.intro,
        Place.image_path,
        Place.type
    ).filter(Place.id == id).first()

    if not place_info:
        return render_template("404.html")

    # 构建当前场所的 node_id
    current_node_id = f"P{int(id):03d}"

    # 从 roadlinks 表获取所有道路连接，构建无向图
    roadlinks = db.session.query(
        RoadLinks.road_id,
        RoadLinks.from_node,
        RoadLinks.to_node,
        RoadLinks.length
    ).all()

    graph = {}
    for link in roadlinks:
        if link.from_node not in graph:
            graph[link.from_node] = []
        if link.to_node not in graph:
            graph[link.to_node] = []
        graph[link.from_node].append((link.to_node, link.length))
        graph[link.to_node].append((link.from_node, link.length))

    # Dijkstra 算法计算最短路径
    def dijkstra(start, target):
        if start not in graph or target not in graph:
            return float('inf')
        dist = {}
        for node in graph:
            dist[node] = float('inf')
        dist[start] = 0
        heap = [(0, start)]

        while heap:
            min_dist = float('inf')
            min_index = -1
            for i in range(len(heap)):
                if heap[i][0] < min_dist:
                    min_dist = heap[i][0]
                    min_index = i
            current_dist, u = heap[min_index]
            del heap[min_index]

            if u == target:
                return current_dist
            if current_dist > dist[u]:
                continue
            for v, w in graph[u]:
                if dist[v] > current_dist + w:
                    dist[v] = current_dist + w
                    heap.append((dist[v], v))
        return dist.get(target, float('inf'))

    # 获取所有服务设施节点（category=1）并计算距离
    service_nodes = db.session.query(Node).filter(Node.category == 1).all()
    nearby_services = []

    for service_node in service_nodes:
        service_id = int(service_node.node_id[1:])  # 从 S001 提取 ID=1
        service = db.session.query(Service).filter(Service.id == service_id).first()
        if not service:
            continue

        # 计算最短距离
        distance = dijkstra(current_node_id, service_node.node_id)
        if distance != float('inf'):
            nearby_services.append({
                'name': service.name,
                'type': service.type,
                'distance': distance,
                'service_id': service_id
            })

    # 类别筛选
    category_filter = request.args.get('category', '')
    search_category = request.args.get('search_category', '').strip()

    if category_filter:
        filtered_services = []
        for service in nearby_services:
            if service['type'] == category_filter:
                filtered_services.append(service)
        nearby_services = filtered_services
    elif search_category:
        filtered_services = []
        for service in nearby_services:
            if search_category in service['type']:
                filtered_services.append(service)
        nearby_services = filtered_services

    # 快速排序实现（升序）
    def quick_sort_ascending(arr):
        if len(arr) <= 1:
            return arr
        pivot = arr[len(arr) // 2]['distance']
        left = []
        middle = []
        right = []
        for service in arr:
            if service['distance'] < pivot:
                left.append(service)
            elif service['distance'] == pivot:
                middle.append(service)
            else:
                right.append(service)
        return quick_sort_ascending(left) + middle + quick_sort_ascending(right)

    # 执行快速排序
    nearby_services = quick_sort_ascending(nearby_services)

    image_path = place_info.image_path or 'images/empty.jpg'
    image_path = image_path.replace('\\', '/')

    return render_template(
        "place3.html",
        place_info=place_info._asdict(),  # 转换为字典方便模板访问
        image_path=image_path,
        nearby_services=nearby_services,
        category=category_filter,
        search_category=search_category,  # 新增传递搜索框输入的类别
        keyword=request.args.get('q', ''),
        sort_by=request.args.get('sort_by', 'recommend')
    )

# app.py 中添加以下路由
@app.route("/place4/<string:id>")
@login_required
def place4(id):
    query_sql = "SELECT * FROM food WHERE id = %s"
    cursor = con_my_sql_with_params(query_sql, (id,))
    food_info = cursor.fetchone()
    if not food_info:
        return render_template("404.html")
    image_path = food_info['image_path'].replace('\\', '/').lstrip('/')
    return render_template("place4.html", food_info=food_info,
                           image_path=image_path)  # 确保模板文件存在

@app.route("/increment_heat/<int:id>")
def increment_heat(id):
    try:
        # 执行更新操作
        update_sql = "UPDATE place SET heat = heat + 1 WHERE id = %s"
        cursor = con_my_sql_with_params(update_sql, (id,))
        cursor.connection.commit()  # 提交事务
        pass
    except Exception as e:
        try:
            cursor.connection.rollback()  # 回滚事务
            print(f"更新热度失败: {e}")
        except Exception as print_error:
            try:
                # 尝试将错误信息写入日志文件
                with open('error.log', 'a', encoding='utf-8') as f:
                    f.write(f"更新热度失败: {repr(e)}\n")
            except Exception as log_error:
                pass
    finally:
        if 'cursor' in locals():
            cursor.close()
    return redirect(url_for('place3', id=id))  # 重定向到place3页面

@app.route("/increment_food_heat/<int:id>")
def increment_food_heat(id):
    try:
        # 执行更新操作
        update_sql = "UPDATE food SET heat = heat + 1 WHERE id = %s"
        cursor = con_my_sql_with_params(update_sql, (id,))
        cursor.connection.commit()  # 提交事务
        print("热度值+1 成功")
    except Exception as e:
        try:
            cursor.connection.rollback()  # 回滚事务
            print(f"更新热度失败: {e}")
        except Exception as print_error:
            try:
                # 尝试将错误信息写入日志文件
                with open('error.log', 'a', encoding='utf-8') as f:
                    f.write(f"更新热度失败: {repr(e)}\n")
            except Exception as log_error:
                pass
    finally:
        if 'cursor' in locals():
            cursor.close()
    return redirect(url_for('place4', id=id))  # 重定向到place4页面

@app.route("/login", methods=["post"])
def login():
    name = request.form.get("username")
    pwd = request.form.get("password")

    user = LoginUser.query.filter_by(username=name).first()
    if user and user.password == pwd:
        session['username'] = name  # 存储用户名到会话中
        return redirect(url_for("index"))
    else:
        return '密码错误 <a href="/">返回登录</a>'


@app.route("/register", methods=["post"])
def register():
    name = request.form.get("username")
    pwd = request.form.get("password")
    pwd2 = request.form.get("password2")
    if pwd2 != pwd:
        return '两次输入密码不一致 <a href="/register">返回注册</a>'

    user = LoginUser.query.filter_by(username=name).first()
    if user:
        return '用户已存在 <a href="/">返回登录</a>'
    else:
        new_user = LoginUser(username=name, password=pwd)
        db.session.add(new_user)
        db.session.commit()
        return '注册成功 <a href="/">返回登录</a>'

@app.route("/logout")
def logout():
    session.pop('username', None)  # 从会话中移除用户名
    return redirect(url_for("index"))

@app.route('/create_diary', methods=['POST'])
def create_diary():
    title = request.form.get('title')
    content = request.form.get('content')
    image = request.files.get('image')
    video = request.files.get('video')
    destination = request.form.get('destination')

    image_path = ""
    if image:
        # 检查 images 目录是否存在，不存在则创建
        if not os.path.exists('images'):
            os.makedirs('images')
        try:
            image_path = "images/" + image.filename
            image.save(image_path)
        except Exception as e:
            return jsonify({"message": f"图片保存失败: {str(e)}"}), 500

    video_path = ""
    if video:
        # 检查 images 目录是否存在，不存在则创建
        if not os.path.exists('images'):
            os.makedirs('images')
        try:
            video_path = "images/" + video.filename
            video.save(video_path)
        except Exception as e:
            return jsonify({"message": f"视频保存失败: {str(e)}"}), 500

    try:
        new_diary = TravelDiary(
            title=title,
            content=content,
            image_path=image_path,
            video_path=video_path,
            destination=destination
        )
        db.session.add(new_diary)
        db.session.commit()
        diary_id = new_diary.id
        add_to_whoosh_index(diary_id, title, content, destination)
        flash('日记创建成功', 'success')  # 添加 flash 消息
        return redirect(url_for('index_cultural'))  # 重定向到 cultural.html 页面
    except Exception as e:
        db.session.rollback()
        return jsonify({"message": f"日记创建失败: {str(e)}"}), 500


@app.route('/search_diaries', methods=['GET'])
def search_diaries():
    search_query = request.args.get('q', '').strip()  # 获取搜索词
    sort_by = request.args.get('sort_by', 'rating')  # 获取排序方式，默认按评分排序

    # 构建基础查询
    base_sql = "SELECT * FROM travel_diaries"

    # 添加搜索条件
    if search_query:
        search_sql = " WHERE title LIKE %s OR content LIKE %s OR destination LIKE %s"
        params = [f'%{search_query}%', f'%{search_query}%', f'%{search_query}%']
    else:
        search_sql = ""
        params = []

    # 添加排序
    if sort_by == 'rating':
        order_sql = " ORDER BY rating DESC"
    elif sort_by == 'views':  # 修正排序条件，使用 views 字段
        order_sql = " ORDER BY views DESC"  # 确保热度高的在前面
    else:
        order_sql = " ORDER BY rating DESC"  # 默认按评分排序

    # 执行查询
    cursor = con_my_sql_with_params(base_sql + search_sql + order_sql, params)
    if isinstance(cursor, tuple):  # 检查是否返回了错误元组
        diaries = []
    else:
        diaries = cursor.fetchall()  # 修改变量名，与模板中使用的变量名一致
    keyword = request.args.get('q', '')
    return render_template('cultural.html', diaries=diaries, keyword=keyword)

@app.route('/get_recommended_diaries', methods=['GET'])
def get_recommended_diaries():
    search_query = request.args.get('q', '').strip()
    sort_by = request.args.get('sort_by', 'rating')
    # 构建基础查询
    base_sql = "SELECT * FROM travel_diaries"
    
    # 添加搜索条件
    search_params = []
    if search_query:
        conditions = []
        conditions.append("title LIKE %s")
        search_params.append(f'%{search_query}%')
        
        conditions.append("content LIKE %s")
        search_params.append(f'%{search_query}%')
        
        conditions.append("destination LIKE %s")
        search_params.append(f'%{search_query}%')
        
        search_sql = " WHERE " + " OR ".join(conditions)
    else:
        search_sql = ""
    
    # 添加排序
    if sort_by == 'views':
        order_by = "views DESC"
    elif sort_by == 'rating':
        order_by = "rating DESC"
    else:
        # 默认按评分排序，防止非法参数
        order_by = "rating DESC"
    
    order_sql = f" ORDER BY {order_by}"
    
    # 执行查询
    try:
        full_query = base_sql + search_sql + order_sql
        print(f"执行SQL查询: {full_query}")
        print(f"查询参数: {search_params}")

        # 使用现有的数据库连接函数
        result = con_my_sql_with_params(full_query, search_params)
        
        # 检查返回结果
        if isinstance(result, tuple):
            # 如果返回的是错误元组
            print(f"数据库查询错误: {result}")
            diaries = []
        else:
            # 假设返回的是游标对象
            diaries = result.fetchall()
            print(f"查询成功，获取到 {len(diaries)} 条记录")
            result.close()  # 关闭游标
    
    except Exception as e:
        print(f"执行查询时出错: {e}")
        diaries = []
    

    # 处理查询结果
    diaries_list = []
    for diary in diaries:
        # 获取评论 (假设评论存储在单独的表中)
        comments = []
        if 'id' in diary:
            # 从数据库获取该日记的评论
            comment_cursor = con_my_sql_with_params(
                "SELECT * FROM comments WHERE diary_id = %s ORDER BY comment_time DESC",
                [diary['id']]
            )
            if comment_cursor and not isinstance(comment_cursor, tuple):
                comments = comment_cursor.fetchall()
                comment_cursor.close()
        
        comment_list = []
        for comment in comments:
            comment_obj = {
                "username": comment.get('username', '匿名用户'),
                "comment": comment.get('comment', ''),
                "comment_time": comment.get('comment_time', '').strftime('%Y-%m-%d %H:%M:%S') 
                if comment.get('comment_time') else ''
            }
            comment_list.append(comment_obj)
        
        diary_dict = {
            "id": diary.get('id'),
            "title": diary.get('title'),
            "content": diary.get('content'),
            "image_path": diary.get('image_path'),
            "video_path": diary.get('video_path'),
            "rating": diary.get('rating', 0),
            "views": diary.get('views', 0),
            "destination": diary.get('destination'),
            "like_count": diary.get('like_count', 0),
            "comments": comment_list  # 添加评论列表
        }
        diaries_list.append(diary_dict)
    
    print(f"返回数据: {len(diaries_list)} 条记录")
    return jsonify(diaries_list)


@app.route('/get_diaries_by_destination', methods=['GET'])
def get_diaries_by_destination():
    destination = request.args.get('destination')
    if destination:
        diaries = TravelDiary.query.filter_by(destination=destination).order_by(TravelDiary.views.desc(), TravelDiary.rating.desc()).all()
    else:
        diaries = TravelDiary.query.all()

    diaries_list = []
    for diary in diaries:
        # 计算点赞数，现在使用 views 字段
        like_count = diary.views  
        # 计算平均评分
        ratings = Rating.query.filter_by(diary_id=diary.id).all()
        if ratings:
            total_rating = sum([r.rating for r in ratings])
            average_rating = total_rating / len(ratings)
        else:
            average_rating = 0

        # 为当前日记获取评论列表
        comment_list = []
        for comment in diary.comments:
            comment_obj = {
                "username": comment.username,
                "comment": comment.comment,
                "comment_time": comment.comment_time.strftime('%Y-%m-%d %H:%M:%S')  # 格式化时间
            }
            comment_list.append(comment_obj)

        diary_dict = {
            "id": diary.id,
            "title": diary.title,
            "content": diary.content,
            "image_path": diary.image_path,
            "video_path": diary.video_path,
            "rating": average_rating,
            "views": diary.views,  # 这里的 views 是总点赞数量
            "destination": diary.destination,
            "like_count": like_count,
            "comments": comment_list  # 添加评论列表
        }
        diaries_list.append(diary_dict)

    return jsonify(diaries_list)

@app.route('/like_diary', methods=['POST'])
def like_diary():
    diary_id = request.form.get('diary_id')
    username = session.get('username')
    if not username:
        return jsonify({"message": "请先登录"}), 401

    try:
        diary_id = int(diary_id)
    except ValueError:
        return jsonify({"message": "无效的日记ID"}), 400
    # 检查日记是否存在
    diary = TravelDiary.query.get(diary_id)
    if not diary:
        return jsonify({"message": "日记不存在"}), 404

    existing_like = Like.query.filter_by(username=username, diary_id=diary_id).first()
    if existing_like:
        return jsonify({"message": "你已经点过赞了"}), 400

    try:
        new_like = Like(username=username, diary_id=diary_id)
        db.session.add(new_like)
        if diary:
            diary.views += 1  # 更新 views 字段，即总点赞数量
        db.session.commit()
        return jsonify({"message": "点赞成功"})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"点赞失败: {str(e)}");
        # 更具体的错误信息返回
        if "UNIQUE constraint failed" in str(e):
            return jsonify({"message": "你已经点过赞了"}), 400
        else:
            return jsonify({"message": f"点赞失败: {str(e)}"}), 500


@app.route('/get_diary_detail/<int:diary_id>', methods=['GET'])
def get_diary_detail(diary_id):
    diary = TravelDiary.query.get(diary_id)
    if not diary:
        return jsonify({"message": "日记不存在"}), 404

    # 计算点赞数，现在使用 views 字段
    like_count = diary.views  
    # 计算平均评分
    ratings = Rating.query.filter_by(diary_id=diary.id).all()
    if ratings:
        total_rating = sum([r.rating for r in ratings])
        average_rating = total_rating / len(ratings)
    else:
        average_rating = 0

    comment_list = []
    for comment in diary.comments:
        comment_obj = {
            "username": comment.username,
            "comment": comment.comment,
            "comment_time": comment.comment_time.strftime('%Y-%m-%d %H:%M:%S')  # 格式化时间
        }
        comment_list.append(comment_obj)

    diary_dict = {
        "id": diary.id,
        "title": diary.title,
        "content": diary.content,
        "image_path": diary.image_path,
        "video_path": diary.video_path,
        "rating": average_rating,
        "views": diary.views,  # 这里的 views 是总点赞数量
        "destination": diary.destination,
        "like_count": like_count,
        "comments": comment_list  # 添加评论列表
    }
    return jsonify(diary_dict)

@app.route('/add_comment', methods=['POST'])
def add_comment():
    diary_id = request.form.get('diary_id')
    username = session.get('username')
    if not username:
        return jsonify({"message": "请先登录"}), 401

    comment = request.form.get('comment')
    if not comment:
        return jsonify({"message": "评论内容不能为空"}), 400

    try:
        diary_id = int(diary_id)
    except ValueError:
        return jsonify({"message": "无效的日记ID"}), 400

    diary = TravelDiary.query.get(diary_id)
    if not diary:
        return jsonify({"message": "日记不存在"}), 404

    new_comment = Comment(
        username=username,
        diary_id=diary_id,
        comment=comment
    )
    db.session.add(new_comment)
    db.session.commit()

    return jsonify({"message": "评论成功"})

@app.route('/rate_diary', methods=['POST'])
def rate_diary():
    diary_id = request.form.get('diary_id')
    username = session.get('username')
    if not username:
        return jsonify({"message": "请先登录"}), 401

    rating = float(request.form.get('rating'))
    if rating < 1 or rating > 10:
        return jsonify({"message": "评分必须在1到10之间"}), 400

    existing_rating = Rating.query.filter_by(username=username, diary_id=diary_id).first()
    if existing_rating:
        existing_rating.rating = rating
    else:
        new_rating = Rating(username=username, diary_id=diary_id, rating=rating)
        db.session.add(new_rating)

    try:
        diary = TravelDiary.query.get(diary_id)
        if diary:
            ratings = Rating.query.filter_by(diary_id=diary.id).all()
            if ratings:
                total_rating = sum([r.rating for r in ratings])
                diary.rating = total_rating / len(ratings)
            else:
                diary.rating = 0
            db.session.commit()
            return jsonify({"message": "评分成功"})
        else:
            return jsonify({"message": "日记不存在"}), 404
    except Exception as e:
        db.session.rollback()
        return jsonify({"message": f"评分失败: {str(e)}"}), 500


@app.route('/compress_diary', methods=['POST'])
def compress_diary():
    diary_id = request.form.get('diary_id')
    diary = TravelDiary.query.get(diary_id)
    if diary:
        # 对日记内容进行压缩
        compressed_content = zlib.compress(diary.content.encode('utf-8'))
        return jsonify({"compressed_content": compressed_content.decode('latin1')})
    else:
        return jsonify({"message": "日记不存在"}), 404

@app.route('/download_diary/<int:diary_id>', methods=['GET'])
def download_diary(diary_id):
    diary = TravelDiary.query.get(diary_id)
    if not diary:
        return jsonify({"message": "日记不存在"}), 404

    download_type = request.args.get('type', 'raw')
    
    # 设置MIME类型和文件扩展名映射
    content_types = {
        'raw': 'text/plain',
        'compressed': 'application/gzip',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    }
    
    # 设置文件扩展名映射
    extensions = {
        'raw': 'txt',
        'compressed': 'txt.gz',
        'docx': 'docx'
    }
    
    try:
        if download_type == 'raw':
            content = f"标题: {diary.title}\n内容: {diary.content}"
            file_obj = BytesIO(content.encode('utf-8'))
            
        elif download_type == 'compressed':
            # 使用gzip进行无损压缩，比zlib更标准
            content = f"标题: {diary.title}\n内容: {diary.content}"
            file_obj = BytesIO()
            with gzip.GzipFile(fileobj=file_obj, mode='wb') as gz:
                gz.write(content.encode('utf-8'))
            file_obj.seek(0)
            
        elif download_type == 'docx':
            doc = Document()
            doc.add_heading(diary.title, 0)
            doc.add_paragraph(diary.content)

            if diary.image_path and os.path.exists(diary.image_path):
                doc.add_heading('图片', level=1)
                try:
                    doc.add_picture(diary.image_path, width=Inches(6))
                except Exception as e:
                    doc.add_paragraph(f"图片插入失败: {str(e)}")

            if diary.video_path:
                doc.add_heading('视频', level=1)
                video_url = f"本地视频路径: {diary.video_path}"
                doc.add_paragraph(video_url)

            file_obj = BytesIO()
            doc.save(file_obj)
            file_obj.seek(0)
            
        else:
            return jsonify({"message": "无效的下载类型"}), 400

        file_name = f"diary_{diary_id}.{extensions[download_type]}"
        
        return send_file(
            file_obj,
            as_attachment=True,
            download_name=file_name,
            mimetype=content_types[download_type]
        )
        
    except Exception as e:
        return jsonify({"message": f"下载失败: {str(e)}"}), 500
    
@app.route('/submit_feedback', methods=['POST'])
@login_required
def submit_feedback():
    mailbox = request.form.get('feedbackEmail')
    content = request.form.get('feedbackMessage')

    if mailbox and content:
        new_feedback = Feedback(mailbox=mailbox, content=content)
        try:
            db.session.add(new_feedback)
            db.session.commit()
            flash('反馈提交成功', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'反馈提交失败: {str(e)}', 'error')
    else:
        flash('邮箱和反馈内容不能为空', 'error')

    return redirect(url_for('index_food'))

@app.route('/api/generate-video', methods=['POST'])
def generate_video_api():
    try:
        # 获取表单数据
        prompt = request.form.get('prompt', '让画面动起来')
        image_file = request.files.get('image')
        
        if not image_file:
            return jsonify({"error": "请上传图片"}), 400

        # 保存上传的图片
        filename = secure_filename(image_file.filename)
        temp_image_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        image_file.save(temp_image_path)

        # 初始化智谱AI客户端
        client = ZhipuAI(api_key=ZHI_PU_API_KEY)

        # 将图片转换为Base64
        def image_to_base64(image_path):
            with open(image_path, "rb") as image_file:
                return base64.b64encode(image_file.read()).decode('utf-8')

        image_base64 = image_to_base64(temp_image_path)

        # 调用视频生成接口
        response = client.videos.generations(
            model="cogvideox-2",
            image_url=image_base64,
            prompt=prompt,
            quality="quality",
            with_audio=True,
            size="1920x1080",
            fps=30,
        )

        # 保存任务信息
        task_id = response.id
        task_status[task_id] = {
            'status': 'PROCESSING',
            'video_url': None,
            'cover_url': None,
            'error': None
        }

        # 启动后台任务检查
        Thread(target=check_video_task_status, args=(task_id,)).start()

        return jsonify({
            "success": True,
            "task_id": task_id,
            "message": "视频生成任务已提交"
        })

    except Exception as e:
        app.logger.error(f"视频生成失败: {str(e)}")
        return jsonify({
            "error": f"视频生成失败: {str(e)}"
        }), 500

@app.route('/api/check-video-status/<task_id>', methods=['GET'])
def check_video_status(task_id):
    status = task_status.get(task_id, {
        'status': 'UNKNOWN',
        'video_url': None,
        'cover_url': None,
        'error': None
    })
    return jsonify(status)

def check_video_task_status(task_id):
    max_attempts = 30
    attempt = 0
    wait_time = 10
    
    while attempt < max_attempts:
        attempt += 1
        try:
            url = f"https://open.bigmodel.cn/api/paas/v4/async-result/{task_id}"
            headers = {
                "Authorization": f"Bearer {ZHI_PU_API_KEY}"
            }
            response = requests.get(url, headers=headers)
            result = response.json()
            
            with status_lock:
                if result.get('task_status') == 'SUCCESS':
                    video_info = result.get('video_result', [{}])[0]
                    task_status[task_id] = {
                        'status': 'SUCCESS',
                        'video_url': video_info.get('url'),
                        'cover_url': video_info.get('cover_image_url'),
                        'error': None
                    }
                    break
                elif result.get('task_status') == 'FAIL':
                    task_status[task_id] = {
                        'status': 'FAIL',
                        'video_url': None,
                        'cover_url': None,
                        'error': result.get('error_message', '生成失败')
                    }
                    break
                else:
                    task_status[task_id]['status'] = result.get('task_status', 'PROCESSING')
            
            time.sleep(wait_time)
            wait_time = min(wait_time * 1.5, 60)  # 最大等待60秒
            
        except Exception as e:
            with status_lock:
                task_status[task_id]['error'] = f"检查状态时出错: {str(e)}"
            time.sleep(wait_time)
    
    if attempt >= max_attempts:
        with status_lock:
            task_status[task_id]['error'] = "检查状态超时"


# A*路径规划算法实现
class AStarPathfinder:
    def __init__(self, transport_mode='walk'):
        self.transport_mode = transport_mode

    def calculate_distance(self, lat1, lon1, lat2, lon2):
        """计算两点间的直线距离（米）"""
        R = 6371000  # 地球半径（米）
        lat1_rad = math.radians(lat1)
        lat2_rad = math.radians(lat2)
        delta_lat = math.radians(lat2 - lat1)
        delta_lon = math.radians(lon2 - lon1)

        a = math.sin(delta_lat/2)**2 + math.cos(lat1_rad) * math.cos(lat2_rad) * math.sin(delta_lon/2)**2
        c = 2 * math.atan2(math.sqrt(a), math.sqrt(1-a))
        distance = R * c
        return distance

    def heuristic(self, node_id, goal_id):
        """启发式函数：计算到目标的估计距离"""
        try:
            node = db.session.query(PathNode).filter(PathNode.id == node_id).first()
            goal = db.session.query(PathNode).filter(PathNode.id == goal_id).first()

            if not node or not goal:
                return float('inf')

            node_location = db.session.query(Location).filter(Location.id == node.location_id).first()
            goal_location = db.session.query(Location).filter(Location.id == goal.location_id).first()

            if not node_location or not goal_location:
                return float('inf')

            return self.calculate_distance(
                node_location.latitude, node_location.longitude,
                goal_location.latitude, goal_location.longitude
            )
        except Exception as e:
            print(f"启发式函数计算错误: {e}")
            return float('inf')

    def get_neighbors(self, node_id):
        """获取节点的邻居节点，根据交通方式过滤合适的路径"""
        try:
            edges = db.session.query(PathEdge).filter(
                db.or_(PathEdge.from_node_id == node_id, PathEdge.to_node_id == node_id)
            ).all()

            neighbors = []
            for edge in edges:
                neighbor_id = edge.to_node_id if edge.from_node_id == node_id else edge.from_node_id

                # 根据交通方式过滤路径类型和计算成本
                if self.transport_mode == 'walk':
                    # 步行：优先选择步行专用道和小径
                    cost = edge.travel_time_walk
                    # 对不同路径类型应用差异化权重
                    if edge.path_type == 'pedestrian':
                        cost *= 0.2   # 步行小径极大优先（捷径）
                    elif edge.path_type == 'walkway':
                        cost *= 0.5   # 步行道大幅优先
                    elif edge.path_type == 'bridge':
                        cost *= 0.7   # 桥梁适合步行
                    elif edge.path_type == 'road':
                        cost *= 2.0   # 公路步行不太安全
                    elif edge.path_type == 'subway':
                        cost *= 2.5   # 地铁需要进出站时间
                    elif edge.path_type == 'bus':
                        cost *= 4.0   # 公交专用道不适合步行
                    elif edge.path_type == 'bike_lane':
                        cost *= 3.0   # 自行车道不太适合步行

                elif self.transport_mode == 'bike':
                    # 骑行：优先选择自行车道和道路
                    cost = edge.travel_time_bike if edge.travel_time_bike > 0 else edge.travel_time_walk * 0.3
                    # 对不同路径类型应用强烈的差异化权重
                    if edge.path_type == 'bike_lane':
                        cost *= 0.1   # 自行车道极大优先
                    elif edge.path_type == 'road':
                        cost *= 0.3   # 道路路径大幅优先
                    elif edge.path_type == 'bridge':
                        cost *= 0.4   # 桥梁路径优先
                    elif edge.path_type == 'walkway':
                        cost *= 2.0   # 步行道中等惩罚
                    elif edge.path_type == 'pedestrian':
                        cost *= 5.0   # 步行小径大幅惩罚（不适合骑行）
                    elif edge.path_type == 'subway':
                        cost *= 20.0  # 地铁路径极大惩罚（骑行不便）
                    elif edge.path_type == 'bus':
                        cost *= 15.0  # 公交专用道大幅惩罚

                elif self.transport_mode == 'bus':
                    # 公交：优先选择公交专用道、地铁和主干道
                    cost = edge.travel_time_bus if edge.travel_time_bus > 0 else edge.travel_time_walk * 2
                    # 对不同路径类型应用强烈的差异化优化
                    if edge.path_type == 'bus':
                        cost *= 0.05  # 公交专用道极大优先
                    elif edge.path_type == 'subway':
                        cost *= 0.1   # 地铁线路极大优先（更快更准时）
                    elif edge.path_type == 'road':
                        cost *= 0.4   # 公路路径大幅优先
                    elif edge.path_type == 'bridge':
                        cost *= 0.6   # 桥梁路径优先
                    elif edge.path_type == 'walkway':
                        cost *= 8.0   # 纯步行道极大惩罚（不适合公交路线）
                    elif edge.path_type == 'pedestrian':
                        cost *= 20.0  # 步行小径极大惩罚
                    elif edge.path_type == 'bike_lane':
                        cost *= 10.0  # 自行车道大幅惩罚
                else:
                    cost = edge.travel_time_walk

                neighbors.append((neighbor_id, cost, edge.distance))

            return neighbors
        except Exception as e:
            app.logger.error(f"获取邻居节点错误: {e}")
            return []

    def find_path(self, start_location_id, goal_location_id, strategy='time'):
        """使用A*算法寻找最优路径"""
        try:
            # 找到起点和终点对应的节点
            start_node = db.session.query(PathNode).filter(PathNode.location_id == start_location_id).first()
            goal_node = db.session.query(PathNode).filter(PathNode.location_id == goal_location_id).first()

            if not start_node or not goal_node:
                return None, "起点或终点不存在"

            start_id = start_node.id
            goal_id = goal_node.id

            # A*算法核心
            open_set = [(0, start_id)]
            closed_set = set()  # 已访问的节点
            came_from = {}
            g_score = {start_id: 0}
            f_score = {start_id: self.heuristic(start_id, goal_id)}

            while open_set:
                current_f, current = heapq.heappop(open_set)

                # 如果已经访问过这个节点，跳过
                if current in closed_set:
                    continue

                # 标记为已访问
                closed_set.add(current)

                if current == goal_id:
                    # 重构路径
                    path = []
                    total_distance = 0
                    total_time = 0

                    while current in came_from:
                        path.append(current)
                        prev = came_from[current]

                        # 计算这一段的距离和时间
                        edge = db.session.query(PathEdge).filter(
                            db.or_(
                                db.and_(PathEdge.from_node_id == prev, PathEdge.to_node_id == current),
                                db.and_(PathEdge.from_node_id == current, PathEdge.to_node_id == prev)
                            )
                        ).first()

                        if edge:
                            total_distance += edge.distance
                            if self.transport_mode == 'walk':
                                total_time += edge.travel_time_walk
                            elif self.transport_mode == 'bike':
                                total_time += edge.travel_time_bike if edge.travel_time_bike > 0 else edge.travel_time_walk * 0.3
                            elif self.transport_mode == 'bus':
                                total_time += edge.travel_time_bus if edge.travel_time_bus > 0 else edge.travel_time_walk * 2

                        current = prev

                    path.append(start_id)
                    path.reverse()

                    return {
                        'path': path,
                        'total_distance': total_distance,
                        'total_time': total_time,
                        'transport_mode': self.transport_mode
                    }, None

                neighbors = self.get_neighbors(current)
                for neighbor, cost, distance in neighbors:
                    # 跳过已访问的邻居节点
                    if neighbor in closed_set:
                        continue

                    tentative_g_score = g_score[current] + (cost if strategy == 'time' else distance)

                    if neighbor not in g_score or tentative_g_score < g_score[neighbor]:
                        came_from[neighbor] = current
                        g_score[neighbor] = tentative_g_score
                        f_score[neighbor] = tentative_g_score + self.heuristic(neighbor, goal_id)
                        heapq.heappush(open_set, (f_score[neighbor], neighbor))

            return None, "无法找到路径"

        except Exception as e:
            return None, f"路径规划错误: {str(e)}"


# 路径规划API接口（临时移除登录要求用于测试）
@app.route('/api/plan_route', methods=['POST'])
def plan_route():
    try:
        data = request.get_json()
        start_location_id = data.get('start_location_id')
        end_location_id = data.get('end_location_id')
        waypoints = data.get('waypoints', [])  # 途经点列表
        transport_mode = data.get('transport_mode', 'walk')
        strategy = data.get('strategy', 'time')  # time 或 distance

        if not start_location_id or not end_location_id:
            return jsonify({'error': '起点和终点不能为空'}), 400

        pathfinder = AStarPathfinder(transport_mode)

        if not waypoints:
            # 简单的点对点路径规划
            result, error = pathfinder.find_path(start_location_id, end_location_id, strategy)
            if error:
                return jsonify({'error': error}), 400

            return jsonify({
                'success': True,
                'route': result,
                'waypoints_count': 0
            })
        else:
            # 多点路径规划
            full_path = []
            total_distance = 0
            total_time = 0

            # 规划从起点到第一个途经点的路径
            current_start = start_location_id
            all_points = waypoints + [end_location_id]

            for waypoint in all_points:
                result, error = pathfinder.find_path(current_start, waypoint, strategy)
                if error:
                    return jsonify({'error': f'无法规划到途经点 {waypoint} 的路径: {error}'}), 400

                if full_path:
                    # 避免重复添加连接点
                    full_path.extend(result['path'][1:])
                else:
                    full_path.extend(result['path'])

                total_distance += result['total_distance']
                total_time += result['total_time']
                current_start = waypoint

            return jsonify({
                'success': True,
                'route': {
                    'path': full_path,
                    'total_distance': total_distance,
                    'total_time': total_time,
                    'transport_mode': transport_mode
                },
                'waypoints_count': len(waypoints)
            })

    except Exception as e:
        return jsonify({'error': f'路径规划失败: {str(e)}'}), 500

# 获取所有地点信息API（临时移除登录要求用于测试）
@app.route('/api/locations', methods=['GET'])
def get_locations():
    try:
        location_type = request.args.get('type', '')
        search_query = request.args.get('q', '')

        query = db.session.query(Location)

        if location_type:
            query = query.filter(Location.type == location_type)

        if search_query:
            query = query.filter(
                db.or_(
                    Location.name.like(f'%{search_query}%'),
                    Location.description.like(f'%{search_query}%')
                )
            )

        locations = query.order_by(Location.heat.desc()).all()

        locations_data = []
        for location in locations:
            locations_data.append({
                'id': location.id,
                'name': location.name,
                'type': location.type,
                'latitude': location.latitude,
                'longitude': location.longitude,
                'description': location.description,
                'image_path': location.image_path,
                'heat': location.heat,
                'score': location.score
            })

        return jsonify({
            'success': True,
            'locations': locations_data,
            'count': len(locations_data)
        })

    except Exception as e:
        return jsonify({'error': f'获取地点信息失败: {str(e)}'}), 500

# 获取路径详细信息API（临时移除登录要求用于测试）
@app.route('/api/route_details', methods=['POST'])
def get_route_details():
    try:
        data = request.get_json()
        path_location_ids = data.get('path', [])

        if not path_location_ids:
            return jsonify({'error': '路径节点不能为空'}), 400

        route_details = []
        for location_id in path_location_ids:
            # 直接通过location_id查找地点信息
            location = db.session.query(Location).filter(Location.id == location_id).first()
            if location:
                # 查找对应的路径节点信息
                node = db.session.query(PathNode).filter(PathNode.location_id == location_id).first()
                route_details.append({
                    'location_id': location.id,
                    'name': location.name,
                    'type': location.type,
                    'latitude': location.latitude,
                    'longitude': location.longitude,
                    'description': location.description,
                    'node_type': node.node_type if node else 'normal'
                })

        return jsonify({
            'success': True,
            'route_details': route_details
        })

    except Exception as e:
        return jsonify({'error': f'获取路径详情失败: {str(e)}'}), 500

# 分析路径类型API
@app.route('/api/analyze_path', methods=['POST'])
def analyze_path():
    try:
        data = request.get_json()
        path_node_ids = data.get('path', [])
        transport_mode = data.get('transport_mode', 'walk')

        if not path_node_ids or len(path_node_ids) < 2:
            return jsonify({'error': '路径节点不足'}), 400

        edge_types = []
        edge_details = []

        # 分析路径中每条边的类型
        for i in range(len(path_node_ids) - 1):
            from_node_id = path_node_ids[i]
            to_node_id = path_node_ids[i + 1]

            # 查找对应的边
            edge = db.session.query(PathEdge).filter(
                db.or_(
                    db.and_(PathEdge.from_node_id == from_node_id, PathEdge.to_node_id == to_node_id),
                    db.and_(PathEdge.from_node_id == to_node_id, PathEdge.to_node_id == from_node_id)
                )
            ).first()

            if edge:
                edge_types.append(edge.path_type)

                # 获取起点和终点的地点信息
                from_location = db.session.query(Location).join(PathNode).filter(PathNode.id == from_node_id).first()
                to_location = db.session.query(Location).join(PathNode).filter(PathNode.id == to_node_id).first()

                edge_details.append({
                    'from': from_location.name if from_location else f'节点{from_node_id}',
                    'to': to_location.name if to_location else f'节点{to_node_id}',
                    'type': edge.path_type,
                    'distance': edge.distance,
                    'travel_time_walk': edge.travel_time_walk,
                    'travel_time_bike': edge.travel_time_bike,
                    'travel_time_bus': edge.travel_time_bus
                })
            else:
                edge_types.append('unknown')
                edge_details.append({
                    'from': f'节点{from_node_id}',
                    'to': f'节点{to_node_id}',
                    'type': 'unknown',
                    'distance': 0,
                    'travel_time_walk': 0,
                    'travel_time_bike': 0,
                    'travel_time_bus': 0
                })

        return jsonify({
            'success': True,
            'edge_types': edge_types,
            'edge_details': edge_details,
            'transport_mode': transport_mode
        })

    except Exception as e:
        return jsonify({'error': f'分析路径失败: {str(e)}'}), 500

# 保存用户路线API
@app.route('/api/save_route', methods=['POST'])
@login_required
def save_route():
    try:
        data = request.get_json()
        route_name = data.get('route_name')
        start_location_id = data.get('start_location_id')
        end_location_id = data.get('end_location_id')
        waypoints = data.get('waypoints', [])
        total_distance = data.get('total_distance')
        total_time = data.get('total_time')
        transport_mode = data.get('transport_mode', 'walk')
        route_data = data.get('route_data', {})

        username = session.get('username')

        if not route_name:
            return jsonify({'error': '路线名称不能为空'}), 400

        new_route = TourRoute(
            user_id=username,
            route_name=route_name,
            start_location_id=start_location_id,
            end_location_id=end_location_id,
            waypoints=json.dumps(waypoints),
            total_distance=total_distance,
            total_time=total_time,
            transport_mode=transport_mode,
            route_data=json.dumps(route_data)
        )

        db.session.add(new_route)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': '路线保存成功',
            'route_id': new_route.id
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'保存路线失败: {str(e)}'}), 500

# 获取用户保存的路线API
@app.route('/api/user_routes', methods=['GET'])
@login_required
def get_user_routes():
    try:
        username = session.get('username')
        routes = db.session.query(TourRoute).filter(TourRoute.user_id == username).order_by(TourRoute.created_at.desc()).all()

        routes_data = []
        for route in routes:
            start_location = db.session.query(Location).filter(Location.id == route.start_location_id).first()
            end_location = db.session.query(Location).filter(Location.id == route.end_location_id).first()

            routes_data.append({
                'id': route.id,
                'route_name': route.route_name,
                'start_location': start_location.name if start_location else '未知',
                'end_location': end_location.name if end_location else '未知',
                'waypoints': json.loads(route.waypoints) if route.waypoints else [],
                'total_distance': route.total_distance,
                'total_time': route.total_time,
                'transport_mode': route.transport_mode,
                'created_at': route.created_at.strftime('%Y-%m-%d %H:%M:%S')
            })

        return jsonify({
            'success': True,
            'routes': routes_data,
            'count': len(routes_data)
        })

    except Exception as e:
        return jsonify({'error': f'获取用户路线失败: {str(e)}'}), 500




if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    create_whoosh_index()
    app.run(debug=True)
    